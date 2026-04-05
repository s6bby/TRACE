from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document

from trace_backend.scanning import DocumentScanner


def test_scan_txt_document(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("CASE SUMMARY\n\nStudent receives reading support.\n\n- Review accommodations", encoding="utf-8")

    result = DocumentScanner().scan_path(path)

    assert result.file_kind == "txt"
    assert result.metrics.page_count == 1
    assert "Student receives reading support." in result.text
    assert any(block.kind == "heading" for block in result.blocks)
    assert any(block.kind == "list_item" for block in result.blocks)


def test_scan_docx_document_extracts_tables(tmp_path: Path) -> None:
    path = tmp_path / "case.docx"
    document = Document()
    document.add_heading("Student Profile", level=1)
    document.add_paragraph("The student requires daily reading intervention.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Goal"
    table.rows[0].cells[1].text = "Support"
    table.rows[1].cells[0].text = "Reading"
    table.rows[1].cells[1].text = "30 minutes daily"
    document.save(path)

    result = DocumentScanner().scan_path(path)

    assert result.file_kind == "docx"
    assert result.metrics.table_count == 1
    assert "Student Profile" in result.text
    assert "30 minutes daily" in result.text
    assert any(block.kind == "heading" for block in result.blocks)
    assert result.tables[0].rows[1][1] == "30 minutes daily"


def test_scan_docx_document_ocrs_inline_images(tmp_path: Path) -> None:
    image_source = fitz.open()
    image_page = image_source.new_page(width=420, height=160)
    image_page.insert_text((30, 90), "IMAGE SUPPORT 45 MINUTES", fontsize=28)
    image_pixmap = image_page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
    image_path = tmp_path / "support.png"
    image_pixmap.save(image_path)
    image_source.close()

    path = tmp_path / "graphic.docx"
    document = Document()
    document.add_heading("Services", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run().add_picture(str(image_path))
    document.save(path)

    result = DocumentScanner().scan_path(path)

    assert result.file_kind == "docx"
    assert result.metrics.image_count == 1
    assert any(block.kind == "image" for block in result.blocks)
    assert "IMAGE SUPPORT" in result.text
    assert any(trace.extractor == "rapidocr" for trace in result.extractor_traces)
    assert any(finding.code == "docx-image-ocr-used" for finding in result.findings)


def test_scan_pdf_document_uses_cross_extractor_verification(tmp_path: Path) -> None:
    path = tmp_path / "case.pdf"
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    page.insert_text((72, 72), "Eligibility Review", fontsize=18)
    page.insert_text(
        (72, 110),
        "The student requires explicit reading support.\nIntervention occurs daily.",
        fontsize=11,
    )

    top = 180
    left = 72
    cell_width = 150
    row_height = 28
    for row in range(3):
        y = top + row * row_height
        page.draw_line((left, y), (left + cell_width * 2, y))
    page.draw_line((left, top + row_height * 3), (left + cell_width * 2, top + row_height * 3))
    for col in range(3):
        x = left + col * cell_width
        page.draw_line((x, top), (x, top + row_height * 3))

    page.insert_text((left + 8, top + 18), "Goal", fontsize=10)
    page.insert_text((left + cell_width + 8, top + 18), "Support", fontsize=10)
    page.insert_text((left + 8, top + row_height + 18), "Reading", fontsize=10)
    page.insert_text((left + cell_width + 8, top + row_height + 18), "Daily", fontsize=10)
    document.save(path)
    document.close()

    result = DocumentScanner().scan_path(path)

    assert result.file_kind == "pdf"
    assert result.metrics.page_count == 1
    assert "Eligibility Review" in result.text
    assert result.verification is not None
    assert len(result.verification.comparisons) == 1
    assert any(table.rows[0][0] == "Goal" for table in result.tables)


def test_scan_pdf_document_uses_ocr_backstop_for_image_heavy_page(tmp_path: Path) -> None:
    image_source = fitz.open()
    image_page = image_source.new_page(width=400, height=200)
    image_page.insert_text((40, 80), "IMAGE ONLY PAGE", fontsize=24)
    image_pixmap = image_page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
    image_source.close()

    pdf_path = tmp_path / "image-only.pdf"
    document = fitz.open()
    page = document.new_page(width=400, height=200)
    page.insert_image(page.rect, pixmap=image_pixmap)
    document.save(pdf_path)
    document.close()

    result = DocumentScanner().scan_path(pdf_path)

    assert result.file_kind == "pdf"
    assert any(trace.extractor == "rapidocr" for trace in result.extractor_traces)
    assert any(finding.code == "pdf-ocr-backstop-used" for finding in result.findings)
    assert "IMAGE ONLY PAGE" in result.text
