from __future__ import annotations

from pathlib import Path

import numpy as np
from docx import Document

from trace_backend.claims.models import Claim, TextSpan
from trace_backend.retrieval import HybridEvidenceRetriever, RetrievalConfig
from trace_backend.scanning import DocumentScanner


class FakeEmbeddingProvider:
    def embed_texts(self, texts: list[str]) -> np.ndarray:
        vectors = []
        for text in texts:
            lowered = text.lower()
            if "decoding" in lowered or "phonics" in lowered:
                vectors.append([1.0, 0.0])
            elif "transport" in lowered or "bus" in lowered:
                vectors.append([0.0, 1.0])
            else:
                vectors.append([0.5, 0.5])
        return np.array(vectors, dtype=float)


def _claim(text: str) -> Claim:
    return Claim(
        claim_id="claim-1",
        text=text,
        source_span=TextSpan(0, len(text), text),
        response_span=text,
    )


def test_retriever_returns_relevant_text_block(tmp_path: Path) -> None:
    path = tmp_path / "case.txt"
    path.write_text(
        "Student Profile\n\nThe student receives daily reading intervention in a small group.",
        encoding="utf-8",
    )
    scan = DocumentScanner().scan_path(path)

    retriever = HybridEvidenceRetriever([scan], config=RetrievalConfig(top_k=3))
    result = retriever.retrieve(_claim("The student receives daily reading intervention."))

    assert result.evidence
    assert "daily reading intervention" in result.evidence[0].snippet.lower()
    assert any(finding.code == "retrieval-semantic-disabled" for finding in result.findings)


def test_retriever_surfaces_table_content(tmp_path: Path) -> None:
    path = tmp_path / "services.docx"
    document = Document()
    document.add_heading("Services", level=1)
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Support"
    table.rows[0].cells[1].text = "Frequency"
    table.rows[1].cells[0].text = "Reading"
    table.rows[1].cells[1].text = "30 minutes daily"
    document.save(path)

    scan = DocumentScanner().scan_path(path)
    retriever = HybridEvidenceRetriever([scan], config=RetrievalConfig(top_k=3))
    result = retriever.retrieve(_claim("Reading support is provided for 30 minutes daily."))

    assert result.evidence
    assert any(item.block_kind == "table" for item in result.evidence)
    assert "30 minutes daily" in result.evidence[0].snippet.lower()


def test_retriever_can_use_embeddings_to_promote_semantic_match(tmp_path: Path) -> None:
    reading_path = tmp_path / "reading.txt"
    reading_path.write_text("Intervention Notes\n\nThe student receives phonics intervention.", encoding="utf-8")
    bus_path = tmp_path / "bus.txt"
    bus_path.write_text("Transportation\n\nThe student rides the district bus home.", encoding="utf-8")

    scans = [DocumentScanner().scan_path(reading_path), DocumentScanner().scan_path(bus_path)]
    retriever = HybridEvidenceRetriever(
        scans,
        config=RetrievalConfig(top_k=2, lexical_weight=0.2, semantic_weight=0.8),
        embedding_provider=FakeEmbeddingProvider(),
    )
    result = retriever.retrieve(_claim("The student receives decoding support."))

    assert result.evidence
    assert result.evidence[0].document_id == "reading"
    assert result.evidence[0].semantic_score is not None
