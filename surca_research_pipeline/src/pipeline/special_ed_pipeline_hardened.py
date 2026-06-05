import argparse
import csv
import hashlib
import json
import platform
import re
import shutil
import sys
import time
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import requests
from docx import Document

REPO_ROOT = Path(__file__).resolve().parents[3]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

try:
    from .scoring_rules import (
        ABSTENTION_PATTERNS,
        FIELD_LABELS,
        FIELD_ORDER,
        PROMPT_FIELD_COVERAGE,
        PROMPT_FILE_MAP,
        PROMPT_LABELS,
        PROMPT_SCORING_NOTES,
        RULES,
    )
except ImportError:
    from scoring_rules import (
        ABSTENTION_PATTERNS,
        FIELD_LABELS,
        FIELD_ORDER,
        PROMPT_FIELD_COVERAGE,
        PROMPT_FILE_MAP,
        PROMPT_LABELS,
        PROMPT_SCORING_NOTES,
        RULES,
    )

try:
    from surca_research_pipeline.src.claims.claim_extractor import build_claim_breakdown
except ImportError:
    from ..claims.claim_extractor import build_claim_breakdown

TEMPERATURE = 0.0
MAX_RESPONSE_TOKENS = 700
DEFAULT_TIMEOUT = 900
DEFAULT_PROMPTS_DIRNAME = "master_prompts"
REPORT_MARKDOWN_NAME = "result_summary.md"
REPORT_JSON_NAME = "summary_metrics.json"
RULE_CATALOG_NAME = "field_rules.json"
DEMO_EXPORT_NAME = "trace-ed-data.json"
TABLE_CELL_SEPARATOR = " | "
TABLE_ROW_SEPARATOR = " || "


def safe_model_name(model_name: str) -> str:
    return model_name.replace("/", "_")


def result_file_stem(case_id: str, model_name: str, prompt_id: str) -> str:
    return f"{case_id}__{safe_model_name(model_name)}__{prompt_id}"

COMMON_MOJIBAKE_REPLACEMENTS = {
    "â€™": "'",
    "â€˜": "'",
    "â€œ": '"',
    "â€\x9d": '"',
    "â€“": "-",
    "â€”": "-",
    "â€¢": "- ",
    "Â ": " ",
    "Â": "",
}

DOCX_NOISE_PATTERNS = [
    r"^authorization for release of records by office of superintendent of public instruction",
    r"^licensed under a ?creative commons attribution",
    r"^\*note: before providing initial special education services",
    r"^purpose: the iep is designed to clearly communicate",
]


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def sha256_file(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as handle:
        while True:
            chunk = handle.read(1024 * 1024)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()


def make_run_id() -> str:
    return "run_" + datetime.utcnow().strftime("%Y%m%d_%H%M%S_%f")


def normalize_run_id(run_id: str) -> str:
    value = run_id.strip()
    if not value:
        raise ValueError("--run-id cannot be empty")
    if not re.fullmatch(r"[A-Za-z0-9._-]+", value):
        raise ValueError("--run-id may only use letters, numbers, dots, dashes, and underscores")
    return value


def dedupe_keep_order(items: List[str]) -> List[str]:
    seen = set()
    result = []
    for item in items:
        cleaned = re.sub(r"\s+", " ", item).strip()
        if cleaned and cleaned not in seen:
            seen.add(cleaned)
            result.append(cleaned)
    return result


def collapse_adjacent_duplicates(items: List[str]) -> List[str]:
    result = []
    for item in items:
        if not item:
            continue
        if result and result[-1] == item:
            continue
        result.append(item)
    return result


def repair_text(text: str) -> str:
    value = text or ""
    value = value.replace("\ufeff", "")

    if any(marker in value for marker in ("â", "Ã", "Â")):
        try:
            repaired = value.encode("latin1").decode("utf-8")
            value = repaired
        except UnicodeError:
            pass

    for broken, fixed in COMMON_MOJIBAKE_REPLACEMENTS.items():
        value = value.replace(broken, fixed)

    value = unicodedata.normalize("NFKC", value)
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    return value


def clean_inline_text(text: str) -> str:
    value = repair_text(text)
    value = value.replace("\x00", " ")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def normalize_for_matching(text: str) -> str:
    value = clean_inline_text(text).lower()
    value = value.replace("–", "-").replace("—", "-")
    return re.sub(r"\s+", " ", value).strip()


def should_skip_docx_chunk(text: str) -> bool:
    normalized = normalize_for_matching(text)
    if not normalized:
        return True
    return any(re.search(pattern, normalized) for pattern in DOCX_NOISE_PATTERNS)


def clean_prompt_text(text: str) -> str:
    value = repair_text(text)
    value = value.replace("•", "- ")
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def load_prompts(base_dir: Path) -> Dict[str, str]:
    prompt_dir = base_dir / DEFAULT_PROMPTS_DIRNAME
    prompts = {}

    for prompt_id, filename in PROMPT_FILE_MAP.items():
        prompt_path = prompt_dir / filename
        if not prompt_path.exists():
            raise FileNotFoundError(f"Missing prompt file: {prompt_path}")

        prompt_text = clean_prompt_text(prompt_path.read_text(encoding="utf-8"))
        if not prompt_text:
            raise ValueError(f"Prompt file is empty: {prompt_path}")

        prompts[prompt_id] = prompt_text

    return prompts


def validate_prompt_files(base_dir: Path) -> List[str]:
    issues = []
    prompt_dir = base_dir / DEFAULT_PROMPTS_DIRNAME

    if not prompt_dir.exists():
        return [f"Missing prompt directory: {prompt_dir}"]

    for prompt_id, filename in PROMPT_FILE_MAP.items():
        prompt_path = prompt_dir / filename

        if not prompt_path.exists():
            issues.append(f"Missing prompt file for {prompt_id}: {filename}")
            continue

        try:
            prompt_text = clean_prompt_text(prompt_path.read_text(encoding="utf-8"))
        except Exception as exc:
            issues.append(f"Could not read prompt file {filename} ({exc})")
            continue

        if not prompt_text:
            issues.append(f"Prompt file is empty: {filename}")
        elif "â" in prompt_text or "Ã" in prompt_text:
            issues.append(f"Prompt file still appears to contain broken encoding: {filename}")

    return issues


def extract_docx_text(path: Path) -> str:
    chunks = []
    doc = Document(str(path))

    for paragraph in doc.paragraphs:
        text = clean_inline_text(paragraph.text)
        if text and not should_skip_docx_chunk(text):
            chunks.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_parts = []

            for cell in row.cells:
                cell_lines = []
                for paragraph in cell.paragraphs:
                    text = clean_inline_text(paragraph.text)
                    if text and not should_skip_docx_chunk(text):
                        cell_lines.append(text)

                cell_lines = dedupe_keep_order(collapse_adjacent_duplicates(cell_lines))
                if cell_lines:
                    row_parts.append(TABLE_CELL_SEPARATOR.join(cell_lines))

            row_parts = dedupe_keep_order(collapse_adjacent_duplicates(row_parts))
            if row_parts:
                chunks.append(TABLE_ROW_SEPARATOR.join(row_parts))

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            text = clean_inline_text(paragraph.text)
            if text and not should_skip_docx_chunk(text):
                chunks.append(text)

        for paragraph in section.footer.paragraphs:
            text = clean_inline_text(paragraph.text)
            if text and not should_skip_docx_chunk(text):
                chunks.append(text)

    final_chunks = dedupe_keep_order(collapse_adjacent_duplicates(chunks))
    return "\n".join(final_chunks)


def flatten_ground_truth(data: Dict[str, Any]) -> Dict[str, bool]:
    if not isinstance(data, dict):
        raise ValueError("ground truth json must contain an object at the top level")

    flat = {}
    for key, value in data.items():
        if key in FIELD_ORDER:
            flat[key] = bool(value)
        elif isinstance(value, dict):
            for subkey, subvalue in value.items():
                if subkey in FIELD_ORDER:
                    flat[subkey] = bool(subvalue)

    missing = [field for field in FIELD_ORDER if field not in flat]
    if missing:
        raise ValueError(f"ground truth missing fields: {missing}")

    return {field: flat[field] for field in FIELD_ORDER}


def find_single_matching_file(case_dir: Path, pattern: str, label: str) -> Path:
    matches = sorted(case_dir.glob(pattern))
    if len(matches) == 0:
        raise FileNotFoundError(f"{case_dir.name} has no {label} file matching {pattern}")
    if len(matches) > 1:
        names = ", ".join(path.name for path in matches)
        raise FileNotFoundError(f"{case_dir.name} has multiple {label} files matching {pattern}: {names}")
    return matches[0]


def find_single_json(case_dir: Path) -> Path:
    return find_single_matching_file(case_dir, "*.json", "json ground truth")


def validate_case_folder(case_dir: Path) -> List[str]:
    issues = []
    iep_files = sorted(case_dir.glob("*IEP.docx"))
    bip_files = sorted(case_dir.glob("*BIP.docx"))
    json_files = sorted(case_dir.glob("*.json"))

    if len(iep_files) != 1:
        if len(iep_files) == 0:
            issues.append(f"{case_dir.name}: missing IEP file matching *IEP.docx")
        else:
            names = ", ".join(path.name for path in iep_files)
            issues.append(f"{case_dir.name}: expected 1 IEP file, found {len(iep_files)} ({names})")

    if len(bip_files) != 1:
        if len(bip_files) == 0:
            issues.append(f"{case_dir.name}: missing BIP file matching *BIP.docx")
        else:
            names = ", ".join(path.name for path in bip_files)
            issues.append(f"{case_dir.name}: expected 1 BIP file, found {len(bip_files)} ({names})")

    if len(json_files) != 1:
        if len(json_files) == 0:
            issues.append(f"{case_dir.name}: missing ground truth json file")
        else:
            names = ", ".join(path.name for path in json_files)
            issues.append(f"{case_dir.name}: expected 1 ground truth json file, found {len(json_files)} ({names})")

    if len(iep_files) == 1:
        try:
            iep_text = extract_docx_text(iep_files[0])
            if not iep_text:
                issues.append(f"{case_dir.name}: IEP text extracted as empty")
        except Exception as exc:
            issues.append(f"{case_dir.name}: could not read IEP docx ({exc})")

    if len(bip_files) == 1:
        try:
            bip_text = extract_docx_text(bip_files[0])
            if not bip_text:
                issues.append(f"{case_dir.name}: BIP text extracted as empty")
        except Exception as exc:
            issues.append(f"{case_dir.name}: could not read BIP docx ({exc})")

    if len(json_files) == 1:
        try:
            raw = json.loads(json_files[0].read_text(encoding="utf-8"))
            flatten_ground_truth(raw)
            expected_case_id = str(raw.get("case_id", "")).strip()
            if expected_case_id and expected_case_id != case_dir.name:
                issues.append(
                    f"{case_dir.name}: ground truth case_id is '{expected_case_id}', expected '{case_dir.name}'"
                )
        except Exception as exc:
            issues.append(f"{case_dir.name}: invalid ground truth json ({exc})")

    return issues


def collect_case_dirs(base_dir: Path, selected_cases: Optional[List[str]]) -> List[Path]:
    case_root = base_dir / "cases"
    if not case_root.exists():
        raise FileNotFoundError(f"Cases directory not found: {case_root}")

    all_case_dirs = sorted(path for path in case_root.iterdir() if path.is_dir())

    if selected_cases:
        requested = dedupe_keep_order(selected_cases)
        case_by_name = {path.name: path for path in all_case_dirs}
        missing = [case_id for case_id in requested if case_id not in case_by_name]
        if missing:
            raise ValueError("Selected case folders were not found: " + ", ".join(missing))
        case_dirs = [case_by_name[case_id] for case_id in requested]
    else:
        case_dirs = all_case_dirs

    if not case_dirs:
        raise ValueError("No matching case folders were found to run.")

    return case_dirs


def format_validation_errors(issues: List[str]) -> str:
    lines = ["Preflight validation failed."]
    lines.extend(f"- {issue}" for issue in issues)
    return "\n".join(lines)


def run_preflight_validation(
    base_dir: Path,
    selected_cases: Optional[List[str]],
    validate_prompts: bool = True,
) -> Dict[str, Any]:
    case_dirs = collect_case_dirs(base_dir, selected_cases)
    issues = []

    if validate_prompts:
        issues.extend(validate_prompt_files(base_dir))

    for case_dir in case_dirs:
        issues.extend(validate_case_folder(case_dir))

    if issues:
        raise ValueError(format_validation_errors(issues))

    return {
        "case_dirs": case_dirs,
        "case_count": len(case_dirs),
    }


def load_case(case_dir: Path) -> Dict[str, Any]:
    iep_path = find_single_matching_file(case_dir, "*IEP.docx", "IEP")
    bip_path = find_single_matching_file(case_dir, "*BIP.docx", "BIP")
    gt_path = find_single_json(case_dir)

    iep_text = extract_docx_text(iep_path)
    bip_text = extract_docx_text(bip_path)
    ground_truth_raw = json.loads(gt_path.read_text(encoding="utf-8"))
    ground_truth = flatten_ground_truth(ground_truth_raw)

    return {
        "case_id": case_dir.name,
        "iep_path": iep_path,
        "bip_path": bip_path,
        "gt_path": gt_path,
        "iep_text": iep_text,
        "bip_text": bip_text,
        "ground_truth": ground_truth,
        "ground_truth_raw": ground_truth_raw,
    }


def get_scored_fields(prompt_id: str) -> List[str]:
    if prompt_id not in PROMPT_FIELD_COVERAGE:
        raise ValueError(f"Unknown prompt id for scoring: {prompt_id}")
    return PROMPT_FIELD_COVERAGE[prompt_id]


def build_model_input(case_id: str, iep_text: str, bip_text: str, prompt_text: str) -> str:
    return f"""You are reviewing synthetic special education documentation for research evaluation.

CASE ID: {case_id}

IEP DOCUMENT:
{iep_text}

BIP DOCUMENT:
{bip_text}

QUESTION:
{prompt_text}

Answer naturally as if responding to a teacher. Do not output JSON unless asked.
""".strip()


def normalize_openai_base(base_url: str) -> str:
    base = base_url.rstrip("/")
    return base if base.endswith("/v1") else base + "/v1"


def safe_error_message(response: requests.Response) -> str:
    try:
        return json.dumps(response.json(), indent=2)
    except Exception:
        return response.text[:4000]


def verify_backend(base_url: str, model_name: str, api_key: Optional[str]) -> None:
    base = normalize_openai_base(base_url)
    headers = {}
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"

    response = requests.get(base + "/models", headers=headers, timeout=30)
    if not response.ok:
        raise RuntimeError(
            f"Backend check failed at {base}/models\n"
            f"Status: {response.status_code}\n"
            f"Body:\n{safe_error_message(response)}"
        )

    model_ids = [item["id"] for item in response.json().get("data", []) if "id" in item]
    if model_name not in model_ids:
        raise ValueError(f"Model '{model_name}' not found on backend.\nAvailable models: {model_ids[:50]}")


def call_openai_compatible(
    base_url: str,
    model_name: str,
    prompt: str,
    api_key: Optional[str],
    temperature: float,
    max_tokens: int,
    timeout: int = DEFAULT_TIMEOUT,
) -> str:
    base = normalize_openai_base(base_url)
    url = base + "/chat/completions"

    headers = {"Content-Type": "application/json"}
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"

    payload = {
        "model": model_name,
        "messages": [
            {
                "role": "system",
                "content": "You are a helpful assistant answering questions about special education documentation.",
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
        "temperature": temperature,
        "max_tokens": max_tokens,
        "stream": False,
    }

    try:
        response = requests.post(url, headers=headers, json=payload, timeout=timeout)
    except requests.RequestException as exc:
        raise RuntimeError(
            f"Request to LM Studio failed.\n"
            f"URL: {url}\n"
            f"Model: {model_name}\n"
            f"Prompt chars: {len(prompt)}"
        ) from exc

    if not response.ok:
        raise RuntimeError(
            f"Chat completion failed at {url}\n"
            f"Status: {response.status_code}\n"
            f"Model: {model_name}\n"
            f"Prompt chars: {len(prompt)}\n"
            f"Max tokens: {max_tokens}\n"
            f"Response body:\n{safe_error_message(response)}"
        )

    data = response.json()
    try:
        content = data["choices"][0]["message"]["content"]
        if isinstance(content, list):
            parts = []
            for item in content:
                if isinstance(item, dict) and item.get("type") == "text":
                    parts.append(item.get("text", ""))
            return "\n".join(parts).strip()
        return str(content).strip()
    except Exception as exc:
        raise RuntimeError("Unexpected response format from backend.\n" f"Body:\n{json.dumps(data, indent=2)}") from exc


def collect_pattern_hits(text: str, patterns: List[str]) -> Tuple[List[str], List[str]]:
    hit_texts = []
    hit_patterns = []

    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            hit_texts.append(match.group(0))
            hit_patterns.append(pattern)

    return dedupe_keep_order(hit_texts)[:5], dedupe_keep_order(hit_patterns)[:5]


def evidence_for_field(text: str, positive: List[str], negative: List[str]) -> Dict[str, List[str]]:
    positive_hits, positive_patterns = collect_pattern_hits(text, positive)
    negative_hits, negative_patterns = collect_pattern_hits(text, negative)
    return {
        "positive_hits": positive_hits,
        "negative_hits": negative_hits,
        "positive_patterns": positive_patterns,
        "negative_patterns": negative_patterns,
    }


def detect_abstention_language(raw_response: str) -> List[Dict[str, str]]:
    text = normalize_for_matching(raw_response)
    hits = []

    for label, pattern in ABSTENTION_PATTERNS.items():
        for match in re.finditer(pattern, text):
            hits.append({"label": label, "text": match.group(0)})

    unique = []
    seen = set()
    for item in hits:
        key = (item["label"], item["text"])
        if key not in seen:
            seen.add(key)
            unique.append(item)

    return unique


def extract_fields_with_evidence(raw_response: str) -> Tuple[Dict[str, bool], Dict[str, Dict[str, List[str]]], List[Dict[str, str]]]:
    text = normalize_for_matching(raw_response)
    fields = {}
    evidence = {}

    for field in FIELD_ORDER:
        rule = RULES[field]
        field_evidence = evidence_for_field(text, rule["positive"], rule["negative"])
        evidence[field] = field_evidence
        if field_evidence["negative_hits"]:
            fields[field] = False
        else:
            fields[field] = bool(field_evidence["positive_hits"])

    abstention_hits = detect_abstention_language(raw_response)
    return fields, evidence, abstention_hits


def compare_to_ground_truth(
    predicted: Dict[str, bool],
    ground_truth: Dict[str, bool],
    scored_fields: List[str],
) -> Tuple[int, int, float, Dict[str, Optional[bool]]]:
    scored_field_set = set(scored_fields)
    matches: Dict[str, Optional[bool]] = {}
    correct = 0

    for field in FIELD_ORDER:
        if field not in scored_field_set:
            matches[field] = None
            continue

        is_match = bool(predicted.get(field, False)) == bool(ground_truth[field])
        matches[field] = is_match
        if is_match:
            correct += 1

    scored_field_count = len(scored_fields)
    accuracy = (correct / scored_field_count) * 100.0 if scored_field_count else 0.0
    return correct, scored_field_count, accuracy, matches


def prepare_run_paths(base_dir: Path, run_id: str, overwrite_run: bool) -> Dict[str, Path]:
    outputs_dir = base_dir / "outputs"
    runs_root = outputs_dir / "runs"
    run_dir = runs_root / run_id

    if run_dir.exists():
        if not overwrite_run:
            raise FileExistsError(
                f"Run directory already exists: {run_dir}\nUse a different --run-id or pass --overwrite-run."
            )
        shutil.rmtree(run_dir)

    paths = {
        "outputs": outputs_dir,
        "runs_root": runs_root,
        "run_dir": run_dir,
        "latest_run": outputs_dir / "latest_run.txt",
        "raw": run_dir / "raw_responses",
        "pred": run_dir / "predicted_json",
        "claims": run_dir / "claims",
        "docx": run_dir / "extracted_case_text",
        "summary_csv": run_dir / "results.csv",
        "field_csv": run_dir / "field_results.csv",
        "manifest": run_dir / "run_manifest.json",
        "report_md": run_dir / REPORT_MARKDOWN_NAME,
        "report_json": run_dir / REPORT_JSON_NAME,
        "rule_catalog": run_dir / RULE_CATALOG_NAME,
    }

    for path in paths.values():
        if path.suffix == "":
            path.mkdir(parents=True, exist_ok=True)

    return paths


def init_csvs(summary_csv: Path, field_csv: Path) -> None:
    if not summary_csv.exists():
        with open(summary_csv, "w", newline="", encoding="utf-8") as handle:
            csv.writer(handle).writerow(
                [
                    "timestamp",
                    "case_id",
                    "model_name",
                    "prompt_id",
                    "prompt_label",
                    "scored_field_count",
                    "correct_scored_fields",
                    "accuracy_percent",
                    "abstention_detected",
                    "abstention_hits",
                    "raw_response_sha256",
                    "predicted_json_sha256",
                ]
            )

    if not field_csv.exists():
        with open(field_csv, "w", newline="", encoding="utf-8") as handle:
            csv.writer(handle).writerow(
                [
                    "timestamp",
                    "case_id",
                    "model_name",
                    "prompt_id",
                    "prompt_label",
                    "field_name",
                    "field_label",
                    "scored_in_prompt",
                    "predicted",
                    "ground_truth",
                    "is_match",
                    "positive_hits",
                    "negative_hits",
                    "matched_positive_patterns",
                    "matched_negative_patterns",
                ]
            )


def write_manifest(manifest_path: Path, info: Dict[str, Any]) -> None:
    manifest_path.write_text(json.dumps(info, indent=2), encoding="utf-8")


def mark_latest_run(paths: Dict[str, Path], run_id: str) -> None:
    paths["latest_run"].write_text(run_id + "\n", encoding="utf-8")


def save_rule_catalog(paths: Dict[str, Path]) -> None:
    payload = {
        "field_order": FIELD_ORDER,
        "field_labels": FIELD_LABELS,
        "rules": RULES,
        "prompt_field_coverage": PROMPT_FIELD_COVERAGE,
        "prompt_scoring_notes": PROMPT_SCORING_NOTES,
        "abstention_patterns": ABSTENTION_PATTERNS,
    }
    paths["rule_catalog"].write_text(json.dumps(payload, indent=2), encoding="utf-8")


def collect_manifest(
    run_id: str,
    mode: str,
    base_dir: Path,
    paths: Dict[str, Path],
    case_dirs: List[Path],
    models: List[str],
    base_url: str,
    prompts: Optional[Dict[str, str]] = None,
) -> Dict[str, Any]:
    prompt_dir = base_dir / DEFAULT_PROMPTS_DIRNAME

    return {
        "run_id": run_id,
        "mode": mode,
        "status": "running",
        "generated_at_utc": datetime.utcnow().isoformat(),
        "base_dir": str(base_dir.resolve()),
        "run_dir": str(paths["run_dir"].resolve()),
        "script_path": str(Path(__file__).resolve()),
        "script_sha256": sha256_file(Path(__file__).resolve()),
        "rules_path": str(paths["rule_catalog"].resolve()),
        "case_ids": [path.name for path in case_dirs],
        "case_count": len(case_dirs),
        "models": models,
        "base_url": normalize_openai_base(base_url) if base_url else "",
        "prompt_dir": str(prompt_dir.resolve()),
        "prompt_files": PROMPT_FILE_MAP,
        "prompts": prompts or {},
        "prompt_labels": PROMPT_LABELS,
        "prompt_field_coverage": PROMPT_FIELD_COVERAGE,
        "prompt_scoring_notes": PROMPT_SCORING_NOTES,
        "field_order": FIELD_ORDER,
        "field_labels": FIELD_LABELS,
        "python_version": platform.python_version(),
        "platform": platform.platform(),
        "temperature": TEMPERATURE,
        "max_tokens": MAX_RESPONSE_TOKENS,
        "notes": [
            "Ground truth is never sent to the tested model.",
            "Prompt text is loaded from the master prompt files.",
            "Scoring is deterministic and regex-based.",
            "Claim extraction splits the response into sentence or bullet-sized units.",
            "Claim filtering keeps only units that look checkable against the source documents.",
            "Prompt-specific scoring follows the testing protocol.",
            "Abstention-style language is recorded at the response level.",
            "Each response is stored before scoring.",
        ],
    }


def save_case_text(paths: Dict[str, Path], case: Dict[str, Any]) -> None:
    (paths["docx"] / f"{case['case_id']}__IEP_extracted.txt").write_text(case["iep_text"], encoding="utf-8")
    (paths["docx"] / f"{case['case_id']}__BIP_extracted.txt").write_text(case["bip_text"], encoding="utf-8")


def read_latest_run_id(base_dir: Path) -> str:
    latest_path = base_dir / "outputs" / "latest_run.txt"
    if latest_path.exists():
        value = latest_path.read_text(encoding="utf-8").strip()
        if value:
            return value

    runs_root = base_dir / "outputs" / "runs"
    if not runs_root.exists():
        raise FileNotFoundError("No runs directory exists yet.")

    run_dirs = sorted([path for path in runs_root.iterdir() if path.is_dir()], key=lambda path: path.name)
    if not run_dirs:
        raise FileNotFoundError("No saved run folders were found.")

    return run_dirs[-1].name


def resolve_run_dir(base_dir: Path, requested_run_id: str) -> Path:
    run_id = normalize_run_id(requested_run_id) if requested_run_id else read_latest_run_id(base_dir)
    run_dir = base_dir / "outputs" / "runs" / run_id
    if not run_dir.exists():
        raise FileNotFoundError(f"Run folder not found: {run_dir}")
    return run_dir


def parse_bool(value: Any) -> bool:
    return str(value).strip().lower() in {"true", "1", "yes"}


def parse_optional_bool(value: Any) -> Optional[bool]:
    lowered = str(value).strip().lower()
    if lowered in {"", "none", "null"}:
        return None
    return lowered in {"true", "1", "yes"}


def load_csv_rows(path: Path) -> List[Dict[str, str]]:
    if not path.exists():
        return []
    with open(path, newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def build_run_metrics(run_dir: Path) -> Dict[str, Any]:
    summary_rows = load_csv_rows(run_dir / "results.csv")
    field_rows = load_csv_rows(run_dir / "field_results.csv")

    if not summary_rows:
        return {
            "run_id": run_dir.name,
            "run_dir": str(run_dir.resolve()),
            "total_evaluations": 0,
            "overall_average_accuracy": 0.0,
            "accuracy_by_prompt": {},
            "accuracy_by_model": {},
            "most_missed_fields": [],
            "top_runs": [],
            "bottom_runs": [],
            "abstention_detected_count": 0,
        }

    def average(values: List[float]) -> float:
        return round(sum(values) / len(values), 2) if values else 0.0

    accuracy_values = [float(row["accuracy_percent"]) for row in summary_rows]
    by_prompt: Dict[str, List[float]] = {}
    by_model: Dict[str, List[float]] = {}

    for row in summary_rows:
        by_prompt.setdefault(row["prompt_id"], []).append(float(row["accuracy_percent"]))
        by_model.setdefault(row["model_name"], []).append(float(row["accuracy_percent"]))

    missed_fields: Dict[str, int] = {}
    for row in field_rows:
        if not parse_bool(row.get("scored_in_prompt", "")):
            continue
        if parse_optional_bool(row.get("is_match", "")) is False:
            field_name = row["field_name"]
            missed_fields[field_name] = missed_fields.get(field_name, 0) + 1

    ranked_rows = sorted(summary_rows, key=lambda row: float(row["accuracy_percent"]), reverse=True)

    metrics = {
        "run_id": run_dir.name,
        "run_dir": str(run_dir.resolve()),
        "case_ids": sorted({row["case_id"] for row in summary_rows}),
        "model_names": sorted({row["model_name"] for row in summary_rows}),
        "prompt_ids": sorted({row["prompt_id"] for row in summary_rows}),
        "total_evaluations": len(summary_rows),
        "overall_average_accuracy": average(accuracy_values),
        "accuracy_by_prompt": {prompt_id: average(values) for prompt_id, values in by_prompt.items()},
        "accuracy_by_model": {model_name: average(values) for model_name, values in by_model.items()},
        "abstention_detected_count": sum(parse_bool(row.get("abstention_detected", "")) for row in summary_rows),
        "most_missed_fields": [
            {
                "field_name": field_name,
                "field_label": FIELD_LABELS.get(field_name, field_name),
                "miss_count": miss_count,
            }
            for field_name, miss_count in sorted(missed_fields.items(), key=lambda item: (-item[1], item[0]))[:10]
        ],
        "top_runs": [
            {
                "case_id": row["case_id"],
                "model_name": row["model_name"],
                "prompt_id": row["prompt_id"],
                "accuracy_percent": float(row["accuracy_percent"]),
                "correct_scored_fields": int(row["correct_scored_fields"]),
                "scored_field_count": int(row["scored_field_count"]),
            }
            for row in ranked_rows[:5]
        ],
        "bottom_runs": [
            {
                "case_id": row["case_id"],
                "model_name": row["model_name"],
                "prompt_id": row["prompt_id"],
                "accuracy_percent": float(row["accuracy_percent"]),
                "correct_scored_fields": int(row["correct_scored_fields"]),
                "scored_field_count": int(row["scored_field_count"]),
            }
            for row in sorted(summary_rows, key=lambda row: float(row["accuracy_percent"]))[:5]
        ],
    }

    return metrics


def render_run_report_markdown(metrics: Dict[str, Any]) -> str:
    lines = [
        f"# TRACE-ED Run Summary: {metrics['run_id']}",
        "",
        f"- Run folder: `{metrics['run_dir']}`",
        f"- Total evaluations: `{metrics['total_evaluations']}`",
        f"- Overall average accuracy: `{metrics['overall_average_accuracy']:.2f}%`",
        f"- Abstention detected in `{metrics['abstention_detected_count']}` response(s)",
        "",
        "## Accuracy by Prompt",
        "",
    ]

    for prompt_id, accuracy in metrics.get("accuracy_by_prompt", {}).items():
        label = PROMPT_LABELS.get(prompt_id, prompt_id)
        lines.append(f"- `{prompt_id}` ({label}): `{accuracy:.2f}%`")

    lines.extend(["", "## Accuracy by Model", ""])

    for model_name, accuracy in metrics.get("accuracy_by_model", {}).items():
        lines.append(f"- `{model_name}`: `{accuracy:.2f}%`")

    lines.extend(["", "## Most Missed Fields", ""])

    if metrics.get("most_missed_fields"):
        for item in metrics["most_missed_fields"]:
            lines.append(f"- `{item['field_label']}`: `{item['miss_count']}` misses")
    else:
        lines.append("- No missed-field data available.")

    lines.extend(["", "## Best Runs", ""])

    for item in metrics.get("top_runs", []):
        lines.append(
            f"- `{item['case_id']}` | `{item['model_name']}` | `{item['prompt_id']}`: "
            f"`{item['correct_scored_fields']}/{item['scored_field_count']}` "
            f"(`{item['accuracy_percent']:.2f}%`)"
        )

    lines.extend(["", "## Weakest Runs", ""])

    for item in metrics.get("bottom_runs", []):
        lines.append(
            f"- `{item['case_id']}` | `{item['model_name']}` | `{item['prompt_id']}`: "
            f"`{item['correct_scored_fields']}/{item['scored_field_count']}` "
            f"(`{item['accuracy_percent']:.2f}%`)"
        )

    return "\n".join(lines).strip() + "\n"


def write_run_report(run_dir: Path) -> Dict[str, Any]:
    metrics = build_run_metrics(run_dir)
    (run_dir / REPORT_JSON_NAME).write_text(json.dumps(metrics, indent=2), encoding="utf-8")
    (run_dir / REPORT_MARKDOWN_NAME).write_text(render_run_report_markdown(metrics), encoding="utf-8")
    return metrics


def load_json_file(path: Path) -> Dict[str, Any]:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def load_or_build_claim_payload(
    run_dir: Path,
    row: Dict[str, str],
    raw_payload: Dict[str, Any],
) -> Dict[str, Any]:
    claim_path = run_dir / "claims" / f"{result_file_stem(row['case_id'], row['model_name'], row['prompt_id'])}.json"
    if claim_path.exists():
        payload = load_json_file(claim_path)
        if payload:
            return payload

    raw_response = raw_payload.get("raw_response", "")
    if not raw_response:
        return {
            "source_unit_count": 0,
            "claim_count": 0,
            "claims_by_type": {},
            "units": [],
            "claims": [],
        }

    payload = {
        "run_id": run_dir.name,
        "timestamp": row["timestamp"],
        "case_id": row["case_id"],
        "model_name": row["model_name"],
        "prompt_id": row["prompt_id"],
        "prompt_label": row.get("prompt_label", PROMPT_LABELS.get(row["prompt_id"], row["prompt_id"])),
        **build_claim_breakdown(raw_response, prefix=result_file_stem(row["case_id"], row["model_name"], row["prompt_id"])),
    }
    claim_path.parent.mkdir(parents=True, exist_ok=True)
    claim_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    return payload


def build_demo_run_payload(run_dir: Path) -> Dict[str, Any]:
    summary_rows = load_csv_rows(run_dir / "results.csv")
    field_rows = load_csv_rows(run_dir / "field_results.csv")
    metrics = load_json_file(run_dir / REPORT_JSON_NAME) or build_run_metrics(run_dir)
    summary_markdown = ""
    summary_md_path = run_dir / REPORT_MARKDOWN_NAME
    if summary_md_path.exists():
        summary_markdown = summary_md_path.read_text(encoding="utf-8")

    field_lookup: Dict[Tuple[str, str, str, str], List[Dict[str, Any]]] = {}
    for row in field_rows:
        key = (row["timestamp"], row["case_id"], row["model_name"], row["prompt_id"])
        field_lookup.setdefault(key, []).append(
            {
                "field_name": row["field_name"],
                "field_label": row.get("field_label", FIELD_LABELS.get(row["field_name"], row["field_name"])),
                "scored_in_prompt": parse_bool(row.get("scored_in_prompt", "")),
                "predicted": parse_bool(row.get("predicted", "")),
                "ground_truth": parse_bool(row.get("ground_truth", "")),
                "is_match": parse_optional_bool(row.get("is_match", "")),
                "positive_hits": row.get("positive_hits", ""),
                "negative_hits": row.get("negative_hits", ""),
                "matched_positive_patterns": row.get("matched_positive_patterns", ""),
                "matched_negative_patterns": row.get("matched_negative_patterns", ""),
            }
        )

    results = []
    for row in summary_rows:
        file_stem = result_file_stem(row["case_id"], row["model_name"], row["prompt_id"])
        raw_path = run_dir / "raw_responses" / f"{file_stem}.json"
        pred_path = run_dir / "predicted_json" / f"{file_stem}.json"
        raw_payload = load_json_file(raw_path)
        pred_payload = load_json_file(pred_path)
        claim_payload = load_or_build_claim_payload(run_dir, row, raw_payload)
        key = (row["timestamp"], row["case_id"], row["model_name"], row["prompt_id"])

        results.append(
            {
                "result_id": f"{file_stem}__{row['timestamp']}",
                "timestamp": row["timestamp"],
                "case_id": row["case_id"],
                "model_name": row["model_name"],
                "prompt_id": row["prompt_id"],
                "prompt_label": row.get("prompt_label", PROMPT_LABELS.get(row["prompt_id"], row["prompt_id"])),
                "accuracy_percent": float(row["accuracy_percent"]),
                "correct_scored_fields": int(row["correct_scored_fields"]),
                "scored_field_count": int(row["scored_field_count"]),
                "abstention_detected": parse_bool(row.get("abstention_detected", "")),
                "abstention_hits": row.get("abstention_hits", ""),
                "raw_response": raw_payload.get("raw_response", ""),
                "prompt_text": raw_payload.get("prompt_text", ""),
                "predicted_fields": pred_payload.get("predicted_fields", {}),
                "matches": pred_payload.get("matches", {}),
                "claim_count": int(claim_payload.get("claim_count", 0)),
                "source_unit_count": int(claim_payload.get("source_unit_count", 0)),
                "claims_by_type": claim_payload.get("claims_by_type", {}),
                "claim_units": claim_payload.get("units", []),
                "claims": claim_payload.get("claims", []),
                "field_results": field_lookup.get(key, []),
            }
        )

    return {
        "run_id": run_dir.name,
        "summary": metrics,
        "summary_markdown": summary_markdown,
        "results": results,
    }


def export_demo_data(base_dir: Path) -> Tuple[Path, int]:
    runs_root = base_dir / "outputs" / "runs"
    run_dirs = []
    if runs_root.exists():
        run_dirs = sorted([path for path in runs_root.iterdir() if path.is_dir()], key=lambda path: path.name)

    payload = {
        "generated_at_utc": datetime.utcnow().isoformat(),
        "runs": [build_demo_run_payload(run_dir) for run_dir in run_dirs if (run_dir / "results.csv").exists()],
    }

    export_path = Path(__file__).resolve().parents[2] / "demo" / "src" / "generated" / DEMO_EXPORT_NAME
    export_path.parent.mkdir(parents=True, exist_ok=True)
    export_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    return export_path, len(payload["runs"])


def format_abstention_hits(hits: List[Dict[str, str]]) -> str:
    if not hits:
        return ""
    parts = [f"{item['label']}:{item['text']}" for item in hits]
    return " | ".join(parts)


def run_extract_only(
    base_dir: Path,
    selected_cases: Optional[List[str]],
    run_id: str,
    overwrite_run: bool,
) -> None:
    preflight = run_preflight_validation(base_dir, selected_cases, validate_prompts=False)
    case_dirs = preflight["case_dirs"]
    paths = prepare_run_paths(base_dir, run_id, overwrite_run)
    manifest = collect_manifest(run_id, "extract", base_dir, paths, case_dirs, [], "", {})
    manifest["preflight_validation_passed"] = True
    manifest["validated_case_count"] = preflight["case_count"]
    manifest["expected_case_count"] = len(case_dirs)
    manifest["completed_case_count"] = 0
    save_rule_catalog(paths)
    write_manifest(paths["manifest"], manifest)
    mark_latest_run(paths, run_id)

    print(f"preflight ok: {preflight['case_count']} case folders ready")
    print(f"run id: {run_id}")
    print(f"run folder: {paths['run_dir']}")

    try:
        for case_dir in case_dirs:
            case = load_case(case_dir)
            save_case_text(paths, case)
            manifest["completed_case_count"] += 1
            print(f"saved extracted text for {case['case_id']}")
    except Exception as exc:
        manifest["status"] = "failed"
        manifest["finished_at_utc"] = datetime.utcnow().isoformat()
        manifest["error"] = str(exc)
        write_manifest(paths["manifest"], manifest)
        raise

    manifest["status"] = "completed"
    manifest["finished_at_utc"] = datetime.utcnow().isoformat()
    write_manifest(paths["manifest"], manifest)


def run_pipeline(
    base_dir: Path,
    models: List[str],
    base_url: str,
    api_key: Optional[str],
    selected_cases: Optional[List[str]],
    run_id: str,
    overwrite_run: bool,
) -> None:
    prompts = load_prompts(base_dir)

    for model_name in models:
        verify_backend(base_url, model_name, api_key)

    preflight = run_preflight_validation(base_dir, selected_cases, validate_prompts=True)
    case_dirs = preflight["case_dirs"]
    paths = prepare_run_paths(base_dir, run_id, overwrite_run)
    init_csvs(paths["summary_csv"], paths["field_csv"])
    save_rule_catalog(paths)

    manifest = collect_manifest(run_id, "run", base_dir, paths, case_dirs, models, base_url, prompts)
    manifest["preflight_validation_passed"] = True
    manifest["validated_case_count"] = preflight["case_count"]
    manifest["expected_run_count"] = len(case_dirs) * len(models) * len(prompts)
    manifest["completed_run_count"] = 0
    write_manifest(paths["manifest"], manifest)
    mark_latest_run(paths, run_id)

    print(f"preflight ok: {preflight['case_count']} case folders ready")
    print(f"run id: {run_id}")
    print(f"run folder: {paths['run_dir']}")

    try:
        for case_dir in case_dirs:
            case = load_case(case_dir)
            save_case_text(paths, case)

            for model_name in models:
                for prompt_id, prompt_text in prompts.items():
                    timestamp = datetime.utcnow().isoformat()
                    scored_fields = get_scored_fields(prompt_id)
                    scored_field_set = set(scored_fields)
                    prompt_label = PROMPT_LABELS.get(prompt_id, prompt_id)

                    model_input = build_model_input(
                        case["case_id"],
                        case["iep_text"],
                        case["bip_text"],
                        prompt_text,
                    )

                    if len(model_input) > 120000:
                        raise RuntimeError(
                            f"{case['case_id']} produced an unusually large prompt "
                            f"({len(model_input)} chars). Check extracted text before sending to LM Studio."
                        )

                    raw_response = call_openai_compatible(
                        base_url=base_url,
                        model_name=model_name,
                        prompt=model_input,
                        api_key=api_key,
                        temperature=TEMPERATURE,
                        max_tokens=MAX_RESPONSE_TOKENS,
                    )

                    predicted, evidence, abstention_hits = extract_fields_with_evidence(raw_response)
                    correct_scored_fields, scored_field_count, accuracy, matches = compare_to_ground_truth(
                        predicted,
                        case["ground_truth"],
                        scored_fields,
                    )

                    file_stem = result_file_stem(case["case_id"], model_name, prompt_id)
                    raw_path = paths["raw"] / f"{file_stem}.json"
                    pred_path = paths["pred"] / f"{file_stem}.json"
                    claim_path = paths["claims"] / f"{file_stem}.json"
                    claim_breakdown = build_claim_breakdown(raw_response, prefix=file_stem)

                    raw_payload = {
                        "run_id": run_id,
                        "timestamp": timestamp,
                        "case_id": case["case_id"],
                        "model_name": model_name,
                        "prompt_id": prompt_id,
                        "prompt_label": prompt_label,
                        "prompt_text": prompt_text,
                        "model_input_sha256": sha256_text(model_input),
                        "raw_response": raw_response,
                    }
                    raw_path.write_text(json.dumps(raw_payload, indent=2), encoding="utf-8")

                    claim_payload = {
                        "run_id": run_id,
                        "timestamp": timestamp,
                        "case_id": case["case_id"],
                        "model_name": model_name,
                        "prompt_id": prompt_id,
                        "prompt_label": prompt_label,
                        **claim_breakdown,
                    }
                    claim_path.write_text(json.dumps(claim_payload, indent=2), encoding="utf-8")

                    pred_payload = {
                        "run_id": run_id,
                        "timestamp": timestamp,
                        "case_id": case["case_id"],
                        "model_name": model_name,
                        "prompt_id": prompt_id,
                        "prompt_label": prompt_label,
                        "scored_fields": scored_fields,
                        "scored_field_count": scored_field_count,
                        "predicted_fields": predicted,
                        "evidence": evidence,
                        "abstention_detected": bool(abstention_hits),
                        "abstention_hits": abstention_hits,
                        "ground_truth_hash": sha256_text(json.dumps(case["ground_truth"], sort_keys=True)),
                        "correct_scored_fields": correct_scored_fields,
                        "matches": matches,
                        "claim_count": claim_breakdown["claim_count"],
                        "claims_by_type": claim_breakdown["claims_by_type"],
                        "accuracy_percent": round(accuracy, 2),
                    }
                    pred_path.write_text(json.dumps(pred_payload, indent=2), encoding="utf-8")

                    with open(paths["summary_csv"], "a", newline="", encoding="utf-8") as handle:
                        csv.writer(handle).writerow(
                            [
                                timestamp,
                                case["case_id"],
                                model_name,
                                prompt_id,
                                prompt_label,
                                scored_field_count,
                                correct_scored_fields,
                                round(accuracy, 2),
                                bool(abstention_hits),
                                format_abstention_hits(abstention_hits),
                                sha256_file(raw_path),
                                sha256_file(pred_path),
                            ]
                        )

                    with open(paths["field_csv"], "a", newline="", encoding="utf-8") as handle:
                        writer = csv.writer(handle)
                        for field in FIELD_ORDER:
                            field_evidence = evidence.get(
                                field,
                                {
                                    "positive_hits": [],
                                    "negative_hits": [],
                                    "positive_patterns": [],
                                    "negative_patterns": [],
                                },
                            )
                            writer.writerow(
                                [
                                    timestamp,
                                    case["case_id"],
                                    model_name,
                                    prompt_id,
                                    prompt_label,
                                    field,
                                    FIELD_LABELS.get(field, field),
                                    field in scored_field_set,
                                    bool(predicted.get(field, False)),
                                    bool(case["ground_truth"][field]),
                                    matches[field],
                                    " | ".join(field_evidence["positive_hits"]),
                                    " | ".join(field_evidence["negative_hits"]),
                                    " | ".join(field_evidence["positive_patterns"]),
                                    " | ".join(field_evidence["negative_patterns"]),
                                ]
                            )

                    manifest["completed_run_count"] += 1
                    print(
                        f"{case['case_id']} | {model_name} | {prompt_id} | "
                        f"{correct_scored_fields}/{scored_field_count} | {accuracy:.2f}%"
                    )
                    time.sleep(0.2)
    except Exception as exc:
        manifest["status"] = "failed"
        manifest["finished_at_utc"] = datetime.utcnow().isoformat()
        manifest["error"] = str(exc)
        write_manifest(paths["manifest"], manifest)
        raise

    metrics = write_run_report(paths["run_dir"])
    manifest["status"] = "completed"
    manifest["finished_at_utc"] = datetime.utcnow().isoformat()
    manifest["report_markdown"] = str(paths["report_md"].resolve())
    manifest["report_json"] = str(paths["report_json"].resolve())
    manifest["overall_average_accuracy"] = metrics["overall_average_accuracy"]
    write_manifest(paths["manifest"], manifest)


def run_report_mode(base_dir: Path, requested_run_id: str) -> None:
    run_dir = resolve_run_dir(base_dir, requested_run_id)
    metrics = write_run_report(run_dir)
    print(f"report ready: {run_dir / REPORT_MARKDOWN_NAME}")
    print(f"overall average accuracy: {metrics['overall_average_accuracy']:.2f}%")


def parse_args():
    parser = argparse.ArgumentParser(description="hardened special education llm evaluation pipeline")
    parser.add_argument("--base-dir", type=str, default="study_pipeline")
    parser.add_argument("--models", type=str, default="", help="comma-separated model ids")
    parser.add_argument("--base-url", type=str, default="http://127.0.0.1:1234")
    parser.add_argument("--api-key", type=str, default="")
    parser.add_argument("--cases", type=str, default="")
    parser.add_argument("--run-id", type=str, default="", help="optional name for the run folder")
    parser.add_argument("--overwrite-run", action="store_true", help="overwrite an existing run folder with the same run id")
    parser.add_argument(
        "--mode",
        type=str,
        default="run",
        choices=["run", "extract", "verify", "validate", "report", "export-demo"],
    )
    return parser.parse_args()


def main():
    args = parse_args()
    base_dir = Path(args.base_dir)
    selected_cases = [item.strip() for item in args.cases.split(",") if item.strip()] or None

    if args.mode == "extract":
        run_id = normalize_run_id(args.run_id) if args.run_id else make_run_id()
        run_extract_only(base_dir, selected_cases, run_id, args.overwrite_run)
        return

    if args.mode == "validate":
        preflight = run_preflight_validation(base_dir, selected_cases, validate_prompts=True)
        print(f"validation passed: {preflight['case_count']} case folders ready")
        return

    if args.mode == "report":
        run_report_mode(base_dir, args.run_id)
        return

    if args.mode == "export-demo":
        export_path, run_count = export_demo_data(base_dir)
        print(f"demo export ready: {export_path}")
        print(f"exported runs: {run_count}")
        return

    models = [item.strip() for item in args.models.split(",") if item.strip()]
    if not models:
        raise ValueError("--models is required in run and verify modes")

    if args.mode == "verify":
        for model_name in models:
            verify_backend(args.base_url, model_name, args.api_key or None)
            print(f"backend ok for model: {model_name}")
        return

    run_id = normalize_run_id(args.run_id) if args.run_id else make_run_id()
    run_pipeline(
        base_dir=base_dir,
        models=models,
        base_url=args.base_url,
        api_key=args.api_key or None,
        selected_cases=selected_cases,
        run_id=run_id,
        overwrite_run=args.overwrite_run,
    )


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(f"error: {exc}", file=sys.stderr)
        raise SystemExit(1)
